import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np
import comtypes.client
from gemguide.constants import GEMCONTENTKEYS, GEMADMINKEYS
from gemguide import fromexcel


wdFormatPDF = 17


def buildDocument(input : str) -> Document:
    document = Document()
    document.add_heading('GEM Study Guide', 0)

    excel = fromexcel.FromExcel(input)

    # handle course title
    title = excel.getCourseItem('Course title')
    ht = document.add_heading(title, 1)

    # handle admin details
    adm_table = document.add_table(0, 2)
    for ix, k in enumerate(GEMADMINKEYS):
        val = excel.getCourseItem(k)

        cells = adm_table.add_row().cells
        cells[0].text = k
        cells[1].text = str(val)

    # handle descriptions and content
    for ix, k in enumerate(GEMCONTENTKEYS):
        val = excel.getCourseItem(k)

        h = document.add_heading(k, 1)
        p = document.add_paragraph(str(val))
        
    # time allocation
    document.add_heading('Time allocation in hours per activity')

    table = document.add_table(0, 2)

    total = 0
    for k, v in excel.allocation.items():
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
        total += int(v)

    cells = table.add_row().cells
    para = cells[0].paragraphs[0]
    run = para.add_run('Sum')
    run.bold = True
    para = cells[1].paragraphs[0]
    run = para.add_run(str(total))
    run.bold = True

    # Assessment, examiners
    document.add_heading('Assessment')
    table = document.add_table(0, 2)
    cells = table.add_row().cells
    cells[0].text = 'Examiners'
    cells[1].text = excel.getCourseItem('Examiner(s)')

    # Assessment, test plan
    table = document.add_table(0, 3)
    head = excel.assessment.columns
    cells = table.add_row().cells
    for j, c in enumerate(head):
        if c:
            para = cells[j].paragraphs[0]
            run = para.add_run(str(c))
            run.bold = True

    for r in excel.assessment.values:
        cells = table.add_row().cells
        for j, c in enumerate(r):
            para = cells[j].paragraphs[0]
            run = para.add_run(str(c))
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return document


def convert2docx(inp : str, outp : str):
    doc = buildDocument(inp)

    doc.save(outp)

def convert2pdf(inp : str, outp : str):
    doc = buildDocument(inp)

    word = comtypes.client.CreateObject('Word.Application')
    temp = 'E:/Data/studyguide/temp.docx'
    doc.save(temp)

    docx = word.Documents.Open(temp)
    docx.SaveAs(outp, FileFormat=wdFormatPDF)
    docx.Close()
    word.Quit()

    os.remove(temp)


